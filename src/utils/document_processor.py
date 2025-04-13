from typing import BinaryIO, Dict, List, Any, Union
import PyPDF2
from docx import Document
import io
import mimetypes
import subprocess
import tempfile
import os
import shutil
from .text_splitter import TextChunker
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import logging
from pathlib import Path
import uuid

class DocumentProcessor:
    def __init__(self):
        """Initialize document processor with text chunker."""
        self.text_chunker = TextChunker()
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.epub'}

    def process_document(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Process a document and return its content and metadata."""
        file_extension = Path(file_name).suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return self._process_pdf(file_path, file_name)
            elif file_extension == '.docx':
                return self._process_word(file_path, file_name)
            elif file_extension == '.doc':
                return self._process_old_word(file_path, file_name)
            elif file_extension == '.txt':
                return self._process_text(file_path, file_name)
            elif file_extension == '.epub':
                return self._process_epub(file_path, file_name)
        except Exception as e:
            logging.error(f"Error processing document {file_name}: {str(e)}")
            raise

    def _process_epub(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Process an EPUB file and extract its content and metadata."""
        try:
            book = epub.read_epub(file_path)
            
            # Extract metadata with safe defaults
            metadata = {
                'filename': file_name,
                'file_type': 'epub',
                'document_id': str(uuid.uuid4()),
                'title': file_name,  # default to filename
                'author': 'Unknown',
                'language': 'Unknown'
            }
            
            # Safely extract metadata from the book
            try:
                if book.get_metadata('DC', 'title'):
                    metadata['title'] = str(book.get_metadata('DC', 'title')[0][0])
                if book.get_metadata('DC', 'creator'):
                    metadata['author'] = str(book.get_metadata('DC', 'creator')[0][0])
                if book.get_metadata('DC', 'language'):
                    metadata['language'] = str(book.get_metadata('DC', 'language')[0][0])
            except Exception as e:
                logging.warning(f"Error extracting some metadata from EPUB: {str(e)}")
            
            # Extract content from all chapters
            content = []
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    # Parse HTML content
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()
                    # Get text and normalize whitespace
                    text = soup.get_text(separator=' ').strip()
                    if text:
                        content.append(text)
            
            # Join all content with newlines
            full_text = '\n\n'.join(content)
            
            # Split text into chunks
            chunks = self.text_chunker.split_text(full_text, metadata)
            
            return {
                'chunks': chunks,
                'metadata': metadata
            }
            
        except Exception as e:
            logging.error(f"Error processing EPUB file {file_name}: {str(e)}")
            raise ValueError(f"Error processing EPUB file: {str(e)}")

    def _process_pdf(self, file_path: str, file_name: str) -> Dict[str, str]:
        """Process PDF files and extract text."""
        try:
            pdf_reader = PyPDF2.PdfReader(file_path)
            text_content = []
            metadata = {
                "title": pdf_reader.metadata.get("/Title", ""),
                "author": pdf_reader.metadata.get("/Author", ""),
                "subject": pdf_reader.metadata.get("/Subject", ""),
                "creator": pdf_reader.metadata.get("/Creator", ""),
                "page_count": len(pdf_reader.pages)
            }
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
            
            # Join all text content
            full_text = "\n\n".join(text_content)
            
            # Split text into chunks
            chunks = self.text_chunker.split_text(full_text, metadata)
            
            return {
                "chunks": chunks,
                "metadata": metadata
            }
        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")

    def _process_word(self, file_path: str, file_name: str) -> Dict[str, str]:
        """Process Word documents and extract text."""
        try:
            temp_file = io.BytesIO(file_path)
            doc = Document(temp_file)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            metadata = {
                "core_properties": {
                    "author": doc.core_properties.author or "",
                    "title": doc.core_properties.title or "",
                    "created": str(doc.core_properties.created or ""),
                    "modified": str(doc.core_properties.modified or "")
                }
            }
            
            # Join all text content
            full_text = "\n\n".join(text_content)
            
            # Split text into chunks
            chunks = self.text_chunker.split_text(full_text, metadata)
            
            return {
                "chunks": chunks,
                "metadata": metadata
            }
        except Exception as e:
            raise ValueError(f"Error processing Word document: {str(e)}")

    def _process_old_word(self, file_path: str, file_name: str) -> Dict[str, str]:
        """Process old Word (.doc) documents using LibreOffice."""
        try:
            temp_base = os.getenv('TMPDIR', '/tmp/libreoffice')
            os.makedirs(temp_base, mode=0o1777, exist_ok=True)
            temp_dir = tempfile.mkdtemp(dir=temp_base)
            os.chmod(temp_dir, 0o1777)
            
            temp_input = os.path.join(temp_dir, "input.doc")
            
            try:
                with open(temp_input, 'wb') as f:
                    content = file_path
                    f.write(content)
                os.chmod(temp_input, 0o666)
                
                env = os.environ.copy()
                env['HOME'] = '/root'
                env['PATH'] = f"/usr/lib/libreoffice/program:{env.get('PATH', '')}"
                
                result = subprocess.run(
                    ['soffice', '--headless', '--convert-to', 'txt:Text', temp_input, '--outdir', temp_dir],
                    capture_output=True,
                    text=True,
                    check=False,
                    env=env
                )
                
                if result.returncode != 0:
                    raise ValueError(f"LibreOffice conversion failed: {result.stderr}")
                
                txt_path = os.path.join(temp_dir, "input.txt")
                if not os.path.exists(txt_path):
                    raise ValueError(f"Converted text file not found at {txt_path}")
                
                with open(txt_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
                
                if not text_content.strip():
                    raise ValueError("No text could be extracted from the document")
                
                metadata = {
                    "format": "doc",
                    "type": "Word 97-2004",
                    "conversion_method": "libreoffice",
                    "original_size": len(content)
                }
                
                # Split text into chunks
                chunks = self.text_chunker.split_text(text_content, metadata)
                
                return {
                    "chunks": chunks,
                    "metadata": metadata
                }
                
            finally:
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
                    
        except Exception as e:
            raise ValueError(f"Error processing Word 97-2004 document: {str(e)}")

    def _process_text(self, file_path: str, file_name: str) -> Dict[str, str]:
        """Process text files and extract content."""
        try:
            content = file_path.read().decode('utf-8')
            metadata = {}
            
            # Split text into chunks
            chunks = self.text_chunker.split_text(content, metadata)
            
            return {
                "chunks": chunks,
                "metadata": metadata
            }
        except Exception as e:
            raise ValueError(f"Error processing text file: {str(e)}")

    @staticmethod
    def get_supported_extensions() -> list:
        """Return list of supported file extensions."""
        return ['pdf', 'doc', 'docx', 'txt', 'md', 'csv', 'epub'] 